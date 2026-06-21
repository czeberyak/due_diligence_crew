"""
Модуль экспорта отчётов Due Diligence в PDF и DOCX форматы.
Сохраняет таблицы, заголовки, списки и форматирование Markdown.
"""

import re
from pathlib import Path
from datetime import datetime

# Для PDF
import markdown
import weasyprint

# Для DOCX
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def export_to_pdf(markdown_text: str, output_path: str) -> str:
    """
    Конвертирует Markdown в PDF через HTML.
    Сохраняет таблицы, заголовки, списки, жирный/курсив.
    """
    # 1. Markdown → HTML
    html_body = markdown.markdown(
        markdown_text,
        extensions=['tables', 'fenced_code', 'toc', 'nl2br']
    )
    
    # 2. Оборачиваем в полноценный HTML с CSS-стилями
    html_full = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: A4;
                margin: 2cm;
            }}
            body {{
                font-family: 'DejaVu Sans', Arial, sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #333;
            }}
            h1 {{
                font-size: 18pt;
                color: #1f77b4;
                border-bottom: 2px solid #1f77b4;
                padding-bottom: 0.3em;
                margin-top: 1.5em;
            }}
            h2 {{
                font-size: 14pt;
                color: #2c3e50;
                margin-top: 1.2em;
            }}
            h3 {{
                font-size: 12pt;
                color: #34495e;
                margin-top: 1em;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin: 1em 0;
                font-size: 10pt;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #1f77b4;
                color: white;
                font-weight: bold;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            ul, ol {{
                margin: 0.5em 0;
                padding-left: 2em;
            }}
            li {{
                margin: 0.3em 0;
            }}
            strong {{
                font-weight: bold;
                color: #c0392b;
            }}
            em {{
                font-style: italic;
            }}
            code {{
                background-color: #f4f4f4;
                padding: 2px 5px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
                font-size: 9pt;
            }}
            pre {{
                background-color: #f4f4f4;
                padding: 1em;
                border-radius: 5px;
                overflow-x: auto;
            }}
            .header {{
                text-align: center;
                margin-bottom: 2em;
            }}
            .footer {{
                text-align: center;
                font-size: 9pt;
                color: #777;
                margin-top: 3em;
                border-top: 1px solid #ddd;
                padding-top: 1em;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>🏭 Due Diligence Report</h1>
            <p><strong>Generated:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </div>
        {html_body}
        <div class="footer">
            <p>Due Diligence Crew v2.0 | AI-powered investment analysis</p>
        </div>
    </body>
    </html>
    """
    
    # 3. HTML → PDF
    pdf_path = Path(output_path)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    
    html_doc = weasyprint.HTML(string=html_full)
    html_doc.write_pdf(str(pdf_path))
    
    return str(pdf_path)


def export_to_docx(markdown_text: str, output_path: str) -> str:
    """
    Конвертирует Markdown в DOCX с сохранением таблиц и форматирования.
    """
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Парсинг Markdown построчно
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Пропускаем пустые строки
        if not line:
            i += 1
            continue
        
        # Заголовки
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            i += 1
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            i += 1
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            i += 1
        
        # Таблицы (начинаются с |)
        elif line.startswith('|') and '|' in line[1:]:
            # Собираем все строки таблицы
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Парсим таблицу
            if table_lines:
                # Извлекаем заголовки
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                
                # Пропускаем разделительную строку (---|---)
                data_start = 1
                if len(table_lines) > 1 and re.match(r'^\|[\s\-:|]+\|$', table_lines[1]):
                    data_start = 2
                
                # Создаём таблицу в DOCX
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Заголовки
                for j, header in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    cell.text = header
                    # Жирный шрифт для заголовков
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                
                # Данные
                for row_idx in range(data_start, len(table_lines)):
                    row_data = [cell.strip() for cell in table_lines[row_idx].split('|')[1:-1]]
                    row = table.add_row()
                    for j, cell_text in enumerate(row_data):
                        if j < len(headers):
                            row.cells[j].text = cell_text
                
                # Добавляем отступ после таблицы
                doc.add_paragraph()
        
        # Списки
        elif line.startswith('- ') or line.startswith('* '):
            # Собираем все элементы списка
            list_items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                list_items.append(lines[i].strip()[2:])
                i += 1
            
            # Создаём маркированный список
            for item in list_items:
                doc.add_paragraph(item, style='List Bullet')
        
        # Нумерованные списки
        elif re.match(r'^\d+\.\s', line):
            list_items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                match = re.match(r'^\d+\.\s(.+)$', lines[i].strip())
                if match:
                    list_items.append(match.group(1))
                i += 1
            
            for item in list_items:
                doc.add_paragraph(item, style='List Number')
        
        # Обычный параграф
        else:
            # Обрабатываем inline-форматирование (**жирный**, *курсив*)
            paragraph = doc.add_paragraph()
            
            # Разбиваем текст по форматированию
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = paragraph.add_run(part[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(part)
            
            i += 1
    
    # Сохраняем документ
    docx_path = Path(output_path)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    
    return str(docx_path)


def export_report(markdown_text: str, base_path: str, formats: list = ['pdf', 'docx']) -> dict:
    """
    Экспортирует отчёт в указанные форматы.
    
    Args:
        markdown_text: Текст отчёта в Markdown
        base_path: Базовый путь для сохранения (без расширения)
        formats: Список форматов ['pdf', 'docx', 'md']
    
    Returns:
        Словарь с путями к созданным файлам
    """
    results = {}
    base = Path(base_path)
    
    if 'md' in formats:
        md_path = str(base.with_suffix('.md'))
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(markdown_text)
        results['md'] = md_path
    
    if 'pdf' in formats:
        try:
            pdf_path = export_to_pdf(markdown_text, str(base.with_suffix('.pdf')))
            results['pdf'] = pdf_path
        except Exception as e:
            print(f"️ Ошибка экспорта в PDF: {e}")
            results['pdf'] = None
    
    if 'docx' in formats:
        try:
            docx_path = export_to_docx(markdown_text, str(base.with_suffix('.docx')))
            results['docx'] = docx_path
        except Exception as e:
            print(f"⚠️ Ошибка экспорта в DOCX: {e}")
            results['docx'] = None
    
    return results